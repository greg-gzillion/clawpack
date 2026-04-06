#!/usr/bin/env python3
"""
DOCUCLAW - Document Creation Agent
ALL rules and templates come from Webclaw (central reference hub)
"""

import sys
import os
import json
import sqlite3
import urllib.parse
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, List, Optional

# ============================================
# CLAWPACK PATHS
# ============================================
AGENT_DIR = Path(__file__).parent
ROOT_DIR = AGENT_DIR.parent.parent
WEBCLAW_REFS = ROOT_DIR / "agents" / "webclaw" / "references" / "docuclaw"
SHARED_DB = Path.home() / ".claw_memory" / "shared_memory.db"
OUTPUT_DIR = AGENT_DIR / "output"

OUTPUT_DIR.mkdir(exist_ok=True)

# ============================================
# WEBCLAW QUERY FUNCTIONS
# ============================================

def get_court_rules_from_webclaw(jurisdiction: str) -> Dict:
    """Query Webclaw for court rules by jurisdiction"""
    rules_path = WEBCLAW_REFS / "court_rules"
    
    if not rules_path.exists():
        return {"error": f"Court rules not found in Webclaw: {rules_path}"}
    
    # Try exact match first (case insensitive)
    jurisdiction_lower = jurisdiction.lower()
    
    for rule_file in rules_path.glob("*.json"):
        try:
            with open(rule_file, 'r', encoding='utf-8') as f:
                rules = json.load(f)
                
                # Check jurisdiction name
                rule_jurisdiction = rules.get("jurisdiction", "").lower()
                if jurisdiction_lower == rule_jurisdiction or jurisdiction_lower in rule_jurisdiction:
                    return rules
                
                # Check if it's a state-level rule file
                if "default_rules" in rules and jurisdiction_lower in rule_file.stem.lower():
                    return rules
        except:
            pass
    
    return {"error": f"No rules found for {jurisdiction} in Webclaw"}

def list_available_jurisdictions() -> List[str]:
    """List all available court rule jurisdictions in Webclaw"""
    rules_path = WEBCLAW_REFS / "court_rules"
    jurisdictions = []
    
    if rules_path.exists():
        for rule_file in rules_path.glob("*.json"):
            try:
                with open(rule_file, 'r', encoding='utf-8') as f:
                    rules = json.load(f)
                    name = rules.get("jurisdiction", rule_file.stem.replace("_", " ").replace("-", " "))
                    jurisdictions.append(name)
            except:
                jurisdictions.append(rule_file.stem.replace("_", " "))
    
    return sorted(jurisdictions)

def get_template_from_webclaw(category: str, template_name: str = None) -> Dict:
    """Query Webclaw for document templates"""
    templates_path = WEBCLAW_REFS / "templates"
    
    if not templates_path.exists():
        return {"error": f"Templates not found in Webclaw: {templates_path}"}
    
    # If specific template requested
    if template_name:
        # Search all categories
        for category_dir in templates_path.iterdir():
            if category_dir.is_dir():
                template_file = category_dir / f"{template_name}.md"
                if template_file.exists():
                    return {
                        "category": category_dir.name,
                        "name": template_name,
                        "content": template_file.read_text(encoding='utf-8')
                    }
        
        # Try with .txt extension
        for category_dir in templates_path.iterdir():
            if category_dir.is_dir():
                template_file = category_dir / f"{template_name}.txt"
                if template_file.exists():
                    return {
                        "category": category_dir.name,
                        "name": template_name,
                        "content": template_file.read_text(encoding='utf-8')
                    }
    
    # Return list of available templates
    templates = []
    for category_dir in templates_path.iterdir():
        if category_dir.is_dir():
            for template_file in category_dir.glob("*.md"):
                templates.append({
                    "category": category_dir.name,
                    "name": template_file.stem,
                    "path": str(template_file)
                })
            for template_file in category_dir.glob("*.txt"):
                templates.append({
                    "category": category_dir.name,
                    "name": template_file.stem,
                    "path": str(template_file)
                })
    
    return {"templates": templates}

def list_template_categories() -> List[str]:
    """List all template categories in Webclaw"""
    templates_path = WEBCLAW_REFS / "templates"
    categories = []
    
    if templates_path.exists():
        for category_dir in templates_path.iterdir():
            if category_dir.is_dir():
                categories.append(category_dir.name)
    
    return sorted(categories)

# ============================================
# DOCUMENT EXPORT FUNCTIONS
# ============================================

def export_to_txt(content: str, output_path: Path) -> bool:
    """Export to plain text"""
    try:
        output_path.write_text(content, encoding='utf-8')
        return True
    except Exception as e:
        print(f"Error saving: {e}")
        return False

def export_to_markdown(content: str, output_path: Path) -> bool:
    """Export to Markdown"""
    return export_to_txt(content, output_path.with_suffix('.md'))

# Try to import optional libraries
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

def export_to_docx(content: str, output_path: Path) -> bool:
    """Export to Microsoft Word"""
    if not DOCX_AVAILABLE:
        return export_to_txt(content, output_path.with_suffix('.txt'))
    
    try:
        from docx import Document
        from docx.shared import Pt
        doc = Document()
        for line in content.split('\n'):
            if line.strip():
                p = doc.add_paragraph(line)
                p.paragraph_format.line_spacing = 1.5
            else:
                doc.add_paragraph()
        doc.save(str(output_path.with_suffix('.docx')))
        return True
    except:
        return export_to_txt(content, output_path.with_suffix('.txt'))

# ============================================
# DOCUMENT CREATION
# ============================================

def create_document(template_content: str, fields: Dict, format: str = "txt") -> Path:
    """Create a document from template content"""
    
    # Add date if not provided
    if "[DATE]" in template_content and "date" not in fields:
        fields["date"] = datetime.now().strftime("%B %d, %Y")
    
    # Substitute fields
    content = template_content
    for key, value in fields.items():
        placeholder = f"[{key.upper()}]"
        content = content.replace(placeholder, str(value))
    
    # Generate filename
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = f"document_{timestamp}"
    
    # Export
    output_path = OUTPUT_DIR / safe_name
    if format == "docx":
        success = export_to_docx(content, output_path)
        output_path = output_path.with_suffix('.docx')
    elif format == "md":
        success = export_to_markdown(content, output_path)
        output_path = output_path.with_suffix('.md')
    else:
        success = export_to_txt(content, output_path)
        output_path = output_path.with_suffix('.txt')
    
    return output_path if success else None

# ============================================
# SHARED MEMORY
# ============================================
SHARED_DB.parent.mkdir(exist_ok=True)

def init_shared_memory():
    conn = sqlite3.connect(str(SHARED_DB))
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS documents_created
                 (id INTEGER PRIMARY KEY, doc_type TEXT, 
                  filename TEXT, jurisdiction TEXT, format TEXT,
                  timestamp TEXT, source_agent TEXT)''')
    conn.commit()
    conn.close()

def save_to_shared_memory(doc_type, filename, jurisdiction, format):
    try:
        conn = sqlite3.connect(str(SHARED_DB))
        c = conn.cursor()
        c.execute('''INSERT INTO documents_created (doc_type, filename, jurisdiction, format, timestamp, source_agent)
                     VALUES (?,?,?,?,?,?)''',
                  (doc_type, filename, jurisdiction, format, datetime.now().isoformat(), "docuclaw"))
        conn.commit()
        conn.close()
        return True
    except:
        return False

# ============================================
# MAIN AGENT CLASS
# ============================================

class DocuClaw:
    def __init__(self):
        init_shared_memory()
        self.print_welcome()
    
    def print_welcome(self):
        print("\n" + "="*70)
        print("📄 DOCUCLAW - Document Creation Agent")
        print("="*70)
        print("ALL RULES & TEMPLATES FROM WEBCLAW (Central Hub)")
        print("="*70)
        print("\n📚 COMMANDS:")
        print("  /jurisdictions           - List available court jurisdictions")
        print("  /rules [jurisdiction]    - Show court rules for a jurisdiction")
        print("  /templates [category]    - List available templates")
        print("  /categories              - List template categories")
        print("  /doc [template] [fields] - Create document")
        print("  /list                    - List created documents")
        print("  /quit                    - Exit")
        print("="*70)
        print(f"📁 Webclaw: {WEBCLAW_REFS}")
        print(f"📁 Output: {OUTPUT_DIR}")
        print("="*70)
    
    def handle_jurisdictions(self):
        jurisdictions = list_available_jurisdictions()
        print("\n🏛️ Available Jurisdictions in Webclaw:")
        for j in jurisdictions:
            print(f"  • {j}")
    
    def handle_rules(self, jurisdiction):
        print(f"\n🏛️ Court Rules for: {jurisdiction}")
        print("-" * 50)
        rules = get_court_rules_from_webclaw(jurisdiction)
        if "error" in rules:
            print(f"❌ {rules['error']}")
        else:
            print(json.dumps(rules, indent=2))
    
    def handle_categories(self):
        categories = list_template_categories()
        print("\n📁 Template Categories in Webclaw:")
        for cat in categories:
            print(f"  • {cat}/")
    
    def handle_templates(self, category=None):
        print("\n📋 Available Templates in Webclaw:")
        result = get_template_from_webclaw(category)
        
        if "templates" in result:
            # Group by category
            by_category = {}
            for t in result["templates"]:
                cat = t["category"]
                if cat not in by_category:
                    by_category[cat] = []
                by_category[cat].append(t["name"])
            
            for cat, templates in by_category.items():
                print(f"\n  📁 {cat}/:")
                for t in templates:
                    print(f"      • {t}")
        elif "error" in result:
            print(f"❌ {result['error']}")
    
    def handle_doc(self, args):
        """Create document: /doc template_name field1=value1 field2=value2 format=docx jurisdiction=Colorado"""
        parts = args.split()
        if not parts:
            print("Usage: /doc [template_name] [field=value ...] [format=docx|txt|md] [jurisdiction=CO]")
            return
        
        template_name = parts[0]
        
        # Parse fields
        fields = {}
        output_format = "txt"
        jurisdiction = None
        
        for part in parts[1:]:
            if "=" in part:
                key, value = part.split("=", 1)
                if key == "format":
                    output_format = value.lower()
                elif key == "jurisdiction":
                    jurisdiction = value
                else:
                    fields[key.lower()] = value
        
        # Get template from Webclaw
        template_result = get_template_from_webclaw(None, template_name)
        
        if "error" in template_result:
            print(f"❌ Template not found: {template_name}")
            print("Use /templates to see available templates")
            return
        
        template_content = template_result.get("content", "")
        
        # Get court rules if jurisdiction provided
        if jurisdiction:
            rules = get_court_rules_from_webclaw(jurisdiction)
            if "error" not in rules:
                fields["court_rules"] = json.dumps(rules.get("formatting_rules", {}), indent=2)
                print(f"✅ Applied {jurisdiction} court rules")
        
        # Create document
        filepath = create_document(template_content, fields, output_format)
        
        if filepath:
            print(f"\n✅ Document created: {filepath}")
            print(f"📄 Template: {template_name}")
            print(f"🎨 Format: {output_format.upper()}")
            if jurisdiction:
                print(f"🏛️ Jurisdiction: {jurisdiction}")
            print(f"📁 Location: {OUTPUT_DIR}")
            
            # Save to shared memory
            save_to_shared_memory(template_name, filepath.name, jurisdiction or "unknown", output_format)
        else:
            print(f"\n❌ Failed to create document")
    
    def handle_list(self):
        files = list(OUTPUT_DIR.glob("*"))
        if files:
            print(f"\n📁 Created documents ({len(files)}):")
            for f in sorted(files, key=lambda x: x.stat().st_mtime, reverse=True)[:10]:
                size = f.stat().st_size
                print(f"  • {f.name} ({size} bytes)")
        else:
            print("\nNo documents created yet")
    
    def run(self):
        self.print_welcome()
        while True:
            try:
                cmd = input("\n📄 DocuClaw> ").strip()
                if not cmd:
                    continue
                
                if cmd == "/quit":
                    print("Goodbye!")
                    break
                elif cmd == "/help":
                    self.print_welcome()
                elif cmd == "/jurisdictions":
                    self.handle_jurisdictions()
                elif cmd == "/categories":
                    self.handle_categories()
                elif cmd.startswith("/rules "):
                    self.handle_rules(cmd[7:])
                elif cmd.startswith("/templates"):
                    parts = cmd.split()
                    if len(parts) > 1:
                        self.handle_templates(parts[1])
                    else:
                        self.handle_templates()
                elif cmd.startswith("/doc "):
                    self.handle_doc(cmd[5:])
                elif cmd == "/list":
                    self.handle_list()
                else:
                    print("Unknown command. Try /help")
                    
            except KeyboardInterrupt:
                print("\nGoodbye!")
                break
            except Exception as e:
                print(f"Error: {e}")

if __name__ == "__main__":
    agent = DocuClaw()
    agent.run()